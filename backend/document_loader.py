"""
Document Loader — supports PDF, DOCX, TXT, and Markdown.
Returns a list of LangChain Document objects with rich metadata.
"""

import hashlib
import logging
import mimetypes
import re
from pathlib import Path
from typing import List

from langchain_core.documents import Document

logger = logging.getLogger(__name__)


# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def _file_hash(path: Path) -> str:
    """SHA-256 hash of a file (first 64 KB for speed)."""
    h = hashlib.sha256()
    with open(path, "rb") as f:
        h.update(f.read(65_536))
    return h.hexdigest()[:16]


def _clean_text(text: str) -> str:
    """Normalise whitespace and remove non-printable characters."""
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r" {2,}", " ", text)
    return text.strip()


# ─────────────────────────────────────────────────────────────────────────────
# Per-format loaders
# ─────────────────────────────────────────────────────────────────────────────

def _load_pdf(path: Path) -> List[Document]:
    try:
        from pypdf import PdfReader
    except ImportError:
        raise ImportError("pypdf is required for PDF loading: pip install pypdf")

    reader = PdfReader(str(path))
    docs: List[Document] = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        text = _clean_text(text)
        if len(text) < 20:
            continue
        docs.append(
            Document(
                page_content=text,
                metadata={"page": i + 1, "total_pages": len(reader.pages)},
            )
        )
    return docs


def _load_docx(path: Path) -> List[Document]:
    try:
        from docx import Document as DocxDocument
    except ImportError:
        raise ImportError("python-docx is required: pip install python-docx")

    doc = DocxDocument(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    full_text = "\n\n".join(paragraphs)
    full_text = _clean_text(full_text)
    return [Document(page_content=full_text, metadata={"paragraphs": len(paragraphs)})]


def _load_txt(path: Path) -> List[Document]:
    for enc in ("utf-8", "latin-1", "cp1252"):
        try:
            text = path.read_text(encoding=enc)
            text = _clean_text(text)
            return [Document(page_content=text, metadata={})]
        except UnicodeDecodeError:
            continue
    raise ValueError(f"Could not decode {path.name} with common encodings.")


# ─────────────────────────────────────────────────────────────────────────────
# Public API
# ─────────────────────────────────────────────────────────────────────────────

LOADER_MAP = {
    ".pdf": _load_pdf,
    ".docx": _load_docx,
    ".txt": _load_txt,
    ".md": _load_txt,
}


def load_document(file_path: str | Path, doc_id: str | None = None) -> List[Document]:
    """
    Load a document from disk and return a list of LangChain Documents.

    Each Document carries base metadata:
        doc_id, filename, extension, file_hash, source
    Plus loader-specific metadata (page number, paragraph count, …).

    Args:
        file_path: Path to the uploaded file.
        doc_id:    Optional unique identifier; a hash is generated if omitted.

    Returns:
        List[Document] — one item per page (PDF) or one item for other formats.
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")

    extension = path.suffix.lower()
    if extension not in LOADER_MAP:
        raise ValueError(
            f"Unsupported file type '{extension}'. "
            f"Supported: {', '.join(LOADER_MAP.keys())}"
        )

    logger.info("Loading document: %s", path.name)
    raw_docs = LOADER_MAP[extension](path)

    file_hash = _file_hash(path)
    doc_id = doc_id or file_hash

    base_meta = {
        "doc_id": doc_id,
        "filename": path.name,
        "extension": extension,
        "file_hash": file_hash,
        "source": str(path),
        "total_sections": len(raw_docs),
    }

    enriched: List[Document] = []
    for idx, doc in enumerate(raw_docs):
        merged_meta = {**base_meta, **doc.metadata, "section_index": idx}
        enriched.append(Document(page_content=doc.page_content, metadata=merged_meta))

    total_chars = sum(len(d.page_content) for d in enriched)
    logger.info(
        "Loaded %d section(s) from '%s' — %d total characters.",
        len(enriched),
        path.name,
        total_chars,
    )
    return enriched
